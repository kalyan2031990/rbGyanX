#!/usr/bin/env python3
"""
rbGyanX v1.0 - NTCP Output QA Reporter Module
==============================================
Part of rbGyanX: Radiobiological Analysis Platform
Version: 1.0.0

Reads the output folder (or zip) created by NTCP analysis modules
and generates:
  1) comprehensive_report.docx – human‑readable QA report
  2) qa_summary_tables.xlsx – per‑organ summary metrics + unique patient list

It flags:
- Inflated patient counts (files vs unique PatientID).
- Unrealistic NTCP values: NaNs, const predictions, outside [0,1].
- Low-n / low-event instability.
- Potential ML overfitting / leakage (AUC≥0.90 with n<40 or events<8).
- Traditional model optimism when events<5 but AUC≥0.85.

Usage
-----
python code4_ntcp_output_QA_reporter.py --input <out_dir_or_zip> --report_outdir <outdir>

Author: rbGyanX Team
License: MIT

Examples
--------
python ntcp_output_QA_reporter.py --input enhanced_ntcp_analysis_ml_out.zip --report_outdir QA_results
python ntcp_output_QA_reporter.py --input ./enhanced_ntcp_analysis_ml_out --report_outdir QA_results
"""

import argparse
import os, re, sys, zipfile, math
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union

import numpy as np
import pandas as pd

# Optional deps: if unavailable, the script will print a helpful error.
try:
    from sklearn.metrics import roc_curve, auc, brier_score_loss
except Exception as e:
    print("[WARN] scikit‑learn not available; AUC/Brier will be NaN. Install scikit‑learn to compute metrics.", file=sys.stderr)
    roc_curve = None
    auc = None
    brier_score_loss = None

try:
    from docx import Document
except Exception as e:
    Document = None
    print("[WARN] python‑docx not available; will skip DOCX report. Install python‑docx.", file=sys.stderr)


def unzip_if_needed(input_path: Path, workdir: Path) -> Path:
    """If input is a .zip, extract to workdir and return the root folder; else return input_path."""
    if input_path.is_file() and input_path.suffix.lower() == ".zip":
        with zipfile.ZipFile(input_path, 'r') as z:
            z.extractall(workdir)
        # Heuristic: if the zip contained a single top-level folder, return that; else return workdir
        entries = [p for p in workdir.iterdir()]
        if len(entries) == 1 and entries[0].is_dir():
            return entries[0]
        return workdir
    return input_path


def discover_files(root: Path) -> List[Path]:
    files = []
    for r, d, fs in os.walk(root):
        for f in fs:
            files.append(Path(r) / f)
    return files


def load_table(path: Path) -> Optional[Union[pd.DataFrame, Dict[str, pd.DataFrame]]]:
    """Load CSV or XLSX (all sheets)."""
    try:
        if path.suffix.lower() == ".csv":
            return pd.read_csv(path)
        if path.suffix.lower() == ".xlsx":
            xl = pd.ExcelFile(path)
            return {s: xl.parse(s) for s in xl.sheet_names}
    except Exception as e:
        print(f"[WARN] Failed to load {path}: {e}", file=sys.stderr)
    return None


def likely_results_file(p: Path) -> bool:
    low = p.name.lower()
    if not low.endswith((".csv", ".xlsx")):
        return False
    keys = ["result", "summary", "by_organ", "metrics", "ntcp", "calc"]
    return any(k in low for k in keys)


def is_patient_level_df(df: pd.DataFrame) -> bool:
    cols = [c.strip().lower() for c in df.columns]
    key_cols = {"patient", "patientid", "id", "mrn"}
    outcome_cols = {"observed_toxicity", "toxicity", "event", "grade", "label"}
    pred_cols = [c for c in cols if c.startswith("ntcp_") or c.startswith("ml_") or "lkb" in c or "rs_" in c]
    return (any(k in cols for k in key_cols) and any(o in cols for o in outcome_cols)) or (len(pred_cols) >= 1 and "organ" in cols)


def harmonize(df: pd.DataFrame) -> pd.DataFrame:
    df2 = df.copy()
    # Rename common columns
    rename_map = {}
    for c in df2.columns:
        lc = c.strip().lower()
        if lc in ("patient", "patientid", "ptid", "id"):
            rename_map[c] = "PatientID"
        elif lc == "organ":
            rename_map[c] = "Organ"
        elif lc in ("observed_toxicity", "toxicity", "event", "label"):
            rename_map[c] = "Observed_Toxicity"
        elif lc == "grade":
            rename_map[c] = "Grade"
    if rename_map:
        df2 = df2.rename(columns=rename_map)

    # Standardize types
    if "Organ" in df2.columns:
        df2["Organ"] = df2["Organ"].astype(str).str.strip()

    # Derive Observed_Toxicity if missing: treat Grade>=2 as positive
    if "Observed_Toxicity" not in df2.columns and "Grade" in df2.columns:
        g = pd.to_numeric(df2["Grade"], errors="coerce")
        df2["Observed_Toxicity"] = (g >= 2).astype(float)

    if "Observed_Toxicity" in df2.columns:
        df2["Observed_Toxicity"] = pd.to_numeric(df2["Observed_Toxicity"], errors="coerce")
        df2["Observed_Toxicity"] = df2["Observed_Toxicity"].fillna(0).clip(0, 1)

    # Normalize prediction columns (keep original names too)
    for c in list(df2.columns):
        lc = c.lower()
        if lc in ("lkb_loglogit", "lkb_probit", "rs_poisson"):
            df2.rename(columns={c: f"NTCP_{c}"}, inplace=True)
        if lc in ("ml_ann", "ml_xgboost"):
            # Standardize ML columns as NTCP_ML_ANN / NTCP_ML_XGBoost
            suffix = "ANN" if "ann" in lc else "XGBoost"
            df2.rename(columns={c: f"NTCP_ML_{suffix}"}, inplace=True)

    return df2


def auc_safe(y_true, y_pred) -> float:
    if roc_curve is None or auc is None:
        return float("nan")
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    if len(y_true) < 5 or len(np.unique(y_true)) < 2:
        return float("nan")
    try:
        fpr, tpr, _ = roc_curve(y_true, y_pred)
        return float(auc(fpr, tpr))
    except Exception:
        return float("nan")


def brier_safe(y_true, y_pred) -> float:
    if brier_score_loss is None:
        return float("nan")
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    if len(y_true) < 5:
        return float("nan")
    try:
        return float(brier_score_loss(y_true, y_pred))
    except Exception:
        return float("nan")


def flag_unrealistic(ntcp_vals: pd.Series) -> List[str]:
    ntcp_vals = pd.to_numeric(ntcp_vals, errors="coerce")
    flags = []
    if ntcp_vals.isna().all():
        flags.append("All NTCP are NaN")
    if (ntcp_vals < 0).any() or (ntcp_vals > 1).any():
        flags.append("NTCP outside [0,1] range")
    if ntcp_vals.nunique(dropna=True) <= 1:
        flags.append("No variation in NTCP predictions")
    return flags


def check_ml_model_quality(ntcp_df: pd.DataFrame, output_dir: Path) -> pd.DataFrame:
    """
    QA Check: ML Model Validation
    ===============================
    Checks for:
    - Overfitting (train vs test AUC difference)
    - Cross-validation consistency
    - Data leakage indicators
    - Sample size adequacy
    """
    
    qa_results = []
    
    if ntcp_df is None or ntcp_df.empty:
        qa_results.append({
            'Check': 'ML Models',
            'Status': 'N/A',
            'Message': 'No data available',
            'Recommendation': 'Verify input data'
        })
        return pd.DataFrame(qa_results)
    
    # Get ML model columns
    ml_cols = [col for col in ntcp_df.columns if 'ML' in col or 'ANN' in col or 'XGBoost' in col or 'XGBOOST' in col]
    
    if not ml_cols:
        qa_results.append({
            'Check': 'ML Models',
            'Status': 'N/A',
            'Message': 'No ML models detected',
            'Recommendation': 'Enable --ml_models flag if needed'
        })
        return pd.DataFrame(qa_results)
    
    # Check 1: Sample size adequacy
    if 'Organ' in ntcp_df.columns:
        for organ in ntcp_df['Organ'].unique():
            organ_df = ntcp_df[ntcp_df['Organ'] == organ]
            n_total = len(organ_df)
            n_events = organ_df['Observed_Toxicity'].sum() if 'Observed_Toxicity' in organ_df.columns else 0
            
            # Rule: Need ≥15 events and ≥50 total for reliable ML
            if n_events < 15:
                qa_results.append({
                    'Check': f'{organ} - Sample Size',
                    'Status': 'WARNING',
                    'Message': f'Only {int(n_events)} events (need ≥15 for ML)',
                    'Recommendation': 'ML predictions may be unreliable'
                })
            elif n_total < 50:
                qa_results.append({
                    'Check': f'{organ} - Sample Size',
                    'Status': 'WARNING',
                    'Message': f'Only {n_total} samples (recommend ≥50 for ML)',
                    'Recommendation': 'Consider external validation'
                })
            else:
                qa_results.append({
                    'Check': f'{organ} - Sample Size',
                    'Status': 'PASS',
                    'Message': f'{n_total} samples, {int(n_events)} events',
                    'Recommendation': 'Adequate for ML training'
                })
    
    # Check 2: Overfitting detection
    # Look for Test AUC vs CV AUC in summary file
    summary_file = output_dir / 'enhanced_summary_performance.csv'
    if not summary_file.exists():
        # Try alternative locations
        for alt_path in [output_dir.parent / 'enhanced_summary_performance.csv',
                        output_dir / 'summary_performance.csv',
                        output_dir.parent / 'summary_performance.csv']:
            if alt_path.exists():
                summary_file = alt_path
                break
    
    if summary_file.exists():
        try:
            summary_df = pd.read_csv(summary_file)
            
            for _, row in summary_df.iterrows():
                organ = row.get('Organ', 'Unknown')
                
                # Check ANN
                if 'AUC_ANN_Test' in row and 'AUC_ANN_CV' in row:
                    test_auc = row['AUC_ANN_Test']
                    cv_auc = row['AUC_ANN_CV']
                    if pd.notna(test_auc) and pd.notna(cv_auc):
                        diff = abs(test_auc - cv_auc)
                        
                        if diff > 0.15:  # >15% difference suggests overfitting
                            qa_results.append({
                                'Check': f'{organ} - ANN Overfitting',
                                'Status': 'FAIL',
                                'Message': f'Test AUC ({test_auc:.3f}) differs from CV AUC ({cv_auc:.3f}) by {diff:.3f}',
                                'Recommendation': 'Possible overfitting - reduce model complexity'
                            })
                        elif diff > 0.10:
                            qa_results.append({
                                'Check': f'{organ} - ANN Overfitting',
                                'Status': 'WARNING',
                                'Message': f'Moderate AUC difference: {diff:.3f}',
                                'Recommendation': 'Monitor performance on external data'
                            })
                        else:
                            qa_results.append({
                                'Check': f'{organ} - ANN Overfitting',
                                'Status': 'PASS',
                                'Message': f'AUC difference acceptable: {diff:.3f}',
                                'Recommendation': 'Model generalizes well'
                            })
                
                # Check XGBoost
                if 'AUC_XGBoost_Test' in row and 'AUC_XGBoost_CV' in row:
                    test_auc = row['AUC_XGBoost_Test']
                    cv_auc = row['AUC_XGBoost_CV']
                    if pd.notna(test_auc) and pd.notna(cv_auc):
                        diff = abs(test_auc - cv_auc)
                        
                        if diff > 0.15:
                            qa_results.append({
                                'Check': f'{organ} - XGBoost Overfitting',
                                'Status': 'FAIL',
                                'Message': f'Test AUC ({test_auc:.3f}) differs from CV AUC ({cv_auc:.3f}) by {diff:.3f}',
                                'Recommendation': 'Possible overfitting - reduce max_depth or n_estimators'
                            })
                        elif diff > 0.10:
                            qa_results.append({
                                'Check': f'{organ} - XGBoost Overfitting',
                                'Status': 'WARNING',
                                'Message': f'Moderate AUC difference: {diff:.3f}',
                                'Recommendation': 'Monitor performance on external data'
                            })
                        else:
                            qa_results.append({
                                'Check': f'{organ} - XGBoost Overfitting',
                                'Status': 'PASS',
                                'Message': f'AUC difference acceptable: {diff:.3f}',
                                'Recommendation': 'Model generalizes well'
                            })
        except Exception as e:
            print(f"[WARN] Could not read summary file for ML validation: {e}", file=sys.stderr)
    
    # Check 3: Data leakage indicators
    # Check for suspiciously perfect performance
    if 'Organ' in ntcp_df.columns:
        for organ in ntcp_df['Organ'].unique():
            organ_df = ntcp_df[ntcp_df['Organ'] == organ]
            
            for ml_col in [c for c in ml_cols if c in organ_df.columns]:
                if 'Observed_Toxicity' in organ_df.columns:
                    y_true = pd.to_numeric(organ_df['Observed_Toxicity'], errors='coerce').fillna(0)
                    y_pred = pd.to_numeric(organ_df[ml_col], errors='coerce')
                    
                    valid_mask = ~(y_pred.isna() | y_true.isna())
                    if valid_mask.sum() > 0:
                        auc_val = auc_safe(y_true[valid_mask], y_pred[valid_mask])
                        
                        if not np.isnan(auc_val) and auc_val > 0.99:
                            qa_results.append({
                                'Check': f'{organ} - Data Leakage',
                                'Status': 'FAIL',
                                'Message': f'{ml_col} has suspiciously high AUC: {auc_val:.3f}',
                                'Recommendation': 'Check for data leakage - outcome variable may be in features'
                            })
    
    # Check 4: Cross-validation consistency
    if summary_file.exists():
        try:
            summary_df = pd.read_csv(summary_file)
            
            for _, row in summary_df.iterrows():
                organ = row.get('Organ', 'Unknown')
                
                # Check ANN CV std
                if 'AUC_ANN_CV' in row and 'AUC_ANN_CV_Std' in row:
                    cv_auc = row['AUC_ANN_CV']
                    cv_std = row['AUC_ANN_CV_Std']
                    
                    if pd.notna(cv_std) and pd.notna(cv_auc):
                        if cv_std > 0.15:
                            qa_results.append({
                                'Check': f'{organ} - ANN CV Consistency',
                                'Status': 'WARNING',
                                'Message': f'High CV variance: {cv_std:.3f} (mean: {cv_auc:.3f})',
                                'Recommendation': 'Model unstable - consider ensemble or more data'
                            })
        except Exception:
            pass
    
    return pd.DataFrame(qa_results)


def main():
    ap = argparse.ArgumentParser(description="QA Reporter for NTCP outputs")
    ap.add_argument("--input", required=True, help="Path to output folder or zip produced by enhanced_ntcp_analysis.py / enhanced_ntcp_ml.py")
    ap.add_argument("--report_outdir", required=False, default="QA_results", help="Directory to write the report and tables")
    args = ap.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    report_dir = Path(args.report_outdir).expanduser().resolve()
    report_dir.mkdir(parents=True, exist_ok=True)
    
    # ✓ CRITICAL: Create QA directory
    qa_dir = report_dir / 'qa_report'
    qa_dir.mkdir(parents=True, exist_ok=True)

    # ✓ FIX: Check for required input files
    if not input_path.exists():
        print(f"[X] Error: Input path not found: {input_path}", file=sys.stderr)
        print("[X] Please run previous steps (Steps 1-4) first to generate analysis outputs", file=sys.stderr)
        sys.exit(2)

    # Prepare working directory for zip extraction if needed
    workdir = report_dir / "_unpacked"
    workdir.mkdir(exist_ok=True)

    root = unzip_if_needed(input_path, workdir) if input_path.exists() else None
    if root is None or not root.exists():
        print(f"[X] Error: Could not process input: {input_path}", file=sys.stderr)
        print("[X] Please check that the input path is valid", file=sys.stderr)
        sys.exit(2)

    files = discover_files(root)
    # Focus on likely results
    candidate_paths = [p for p in files if likely_results_file(p)]
    tables = {}
    for p in candidate_paths:
        tables[str(p)] = load_table(p)

    # Gather patient-level frames
    patient_frames: List[Tuple[str, pd.DataFrame]] = []
    for path, obj in tables.items():
        if obj is None:
            continue
        if isinstance(obj, dict):
            for sname, df in obj.items():
                if isinstance(df, pd.DataFrame) and df.shape[0] > 0 and is_patient_level_df(df):
                    patient_frames.append((f"{path}::{sname}", df))
        elif isinstance(obj, pd.DataFrame):
            if obj.shape[0] > 0 and is_patient_level_df(obj):
                patient_frames.append((path, obj))

    # Harmonize and combine
    harm_frames: List[Tuple[str, pd.DataFrame]] = []
    for name, df in patient_frames:
        try:
            hf = harmonize(df)
            if "Organ" in hf.columns and ("Observed_Toxicity" in hf.columns or "Grade" in hf.columns):
                # Ensure Observed_Toxicity exists
                if "Observed_Toxicity" not in hf.columns and "Grade" in hf.columns:
                    g = pd.to_numeric(hf["Grade"], errors="coerce")
                    hf["Observed_Toxicity"] = (g >= 2).astype(float)
                harm_frames.append((name, hf))
        except Exception as e:
            print(f"[WARN] Harmonize failed for {name}: {e}", file=sys.stderr)

    combined = pd.concat([hf for _, hf in harm_frames], ignore_index=True, sort=False) if harm_frames else None

    # Determine unique patients
    def get_pid_cols(df: pd.DataFrame) -> List[str]:
        return [c for c in df.columns if c.lower() in ("patientid", "patient", "id")]
    patient_ids = set()
    if combined is not None:
        pid_cols = get_pid_cols(combined)
        if pid_cols:
            for pid in combined[pid_cols[0]].astype(str).str.strip().unique():
                if pid:
                    patient_ids.add(pid)
        else:
            # fallback via filename patterns: e.g., PID_Organ.csv
            for p in files:
                m = re.search(r"[\\/](\w+)_([A-Za-z]+)\.csv$", str(p))
                if m:
                    patient_ids.add(m.group(1))

    # Compute per-organ metrics and flags
    issues: List[str] = []
    report_rows: List[Dict[str, Union[str, int, float]]] = []
    if combined is not None and "Organ" in combined.columns:
        organs = sorted([o for o in combined["Organ"].dropna().unique()])
        model_cols = [c for c in combined.columns if c.startswith("NTCP_")]
        # Allow some common alt names
        aliases = {"NTCP_LKB_LogLogit":"NTCP_lkb_loglogit", "NTCP_LKB_Probit":"NTCP_lkb_probit", "NTCP_RS_Poisson":"NTCP_rs_poisson"}
        for organ in organs:
            sub = combined[combined["Organ"] == organ].copy()
            n = len(sub)
            if "Observed_Toxicity" in sub.columns:
                events = int(pd.to_numeric(sub["Observed_Toxicity"], errors="coerce").fillna(0).sum())
            else:
                events = np.nan
            event_rate = (events/n*100 if n>0 and not np.isnan(events) else np.nan)

            # Metrics
            perf = {}
            for col in model_cols:
                if col in sub.columns:
                    y_true = pd.to_numeric(sub["Observed_Toxicity"], errors="coerce").fillna(0).values if "Observed_Toxicity" in sub.columns else None
                    y_pred = pd.to_numeric(sub[col], errors="coerce").values
                    perf[col] = {
                        "AUC": auc_safe(y_true, y_pred) if y_true is not None else float("nan"),
                        "Brier": brier_safe(y_true, y_pred) if y_true is not None else float("nan"),
                        "Flags": flag_unrealistic(sub[col])
                    }

            # Data quality flags
            if n < 20:
                issues.append(f"{organ}: Low sample size (n={n})")
            if not np.isnan(events) and events < 5:
                issues.append(f"{organ}: Few events (events={events})")

            # ML overfitting/leakage heuristic
            for ml_key in ["NTCP_ML_ANN", "NTCP_ML_XGBoost", "NTCP_ML_XGBOOST"]:
                if ml_key in perf:
                    auc_v = perf[ml_key]["AUC"]
                    if not (auc_v is None or np.isnan(auc_v)):
                        if (n < 40 or (not np.isnan(events) and events < 8)) and auc_v >= 0.90:
                            issues.append(f"{organ}: Potential ML overfitting/leakage (AUC={auc_v:.3f}, n={n}, events={events})")

            # Traditional models optimism under very low events
            for trad in ["NTCP_LKB_LogLogit", "NTCP_LKB_Probit", "NTCP_RS_Poisson"]:
                if trad in perf:
                    auc_v = perf[trad]["AUC"]
                    if not (auc_v is None or np.isnan(auc_v)):
                        if (not np.isnan(events) and events < 5) and auc_v >= 0.85:
                            issues.append(f"{organ}: Traditional model AUC={auc_v:.3f} with events={events} (unstable)")

            # Tabulate
            best_model, best_auc = None, -1
            for k, v in perf.items():
                if v["AUC"] is not None and not np.isnan(v["AUC"]) and v["AUC"] > best_auc:
                    best_auc = v["AUC"]
                    best_model = k

            row = {
                "Organ": organ, "n": n, "events": events,
                "event_rate_%": (round(event_rate, 1) if isinstance(event_rate, (int, float)) and not np.isnan(event_rate) else np.nan),
                "best_model": best_model, "best_auc": (round(best_auc, 3) if best_auc >= 0 else np.nan),
            }
            for k, v in perf.items():
                row[f"AUC|{k}"] = (round(v["AUC"], 3) if v["AUC"] == v["AUC"] else np.nan)
                row[f"Brier|{k}"] = (round(v["Brier"], 3) if v["Brier"] == v["Brier"] else np.nan)
                row[f"Flags|{k}"] = "; ".join(v["Flags"]) if v["Flags"] else ""
            report_rows.append(row)
    else:
        issues.append("Could not locate a patient-level results table or Organ column in outputs.")

    summary_df = pd.DataFrame(report_rows) if report_rows else pd.DataFrame()

    # Global stats
    global_rows = int(summary_df["n"].sum()) if "n" in summary_df else np.nan
    global_patients = len(patient_ids) if patient_ids else np.nan
    
    # Enhanced validation checks
    validation_checks = []
    
    # Check 1: Patient count consistency (from Step 1)
    if global_patients:
        validation_checks.append({
            'Check': 'Patient Count Consistency',
            'Expected': '57 (from Step 1 preprocessing)',
            'Actual': f'{global_patients}',
            'Status': 'PASS' if global_patients <= 57 else 'REVIEW',
            'Details': f'Found {global_patients} unique patients. Expected ~57 from Step 1.'
        })
    
    # Check 2: Sample size per organ
    if not summary_df.empty and "n" in summary_df.columns:
        for _, row in summary_df.iterrows():
            organ = row.get("Organ", "Unknown")
            n_samples = row.get("n", 0)
            expected_min = 20  # Minimum for reliable ML training
            
            validation_checks.append({
                'Check': f'{organ} Sample Size',
                'Expected': f'>= {expected_min} for ML training',
                'Actual': f'{n_samples}',
                'Status': 'PASS' if n_samples >= expected_min else 'WARNING',
                'Details': f'{organ}: {n_samples} samples. ML training requires >= {expected_min} samples.'
            })
    
    # Check 3: Event rate validation
    if not summary_df.empty and "events" in summary_df.columns and "n" in summary_df.columns:
        for _, row in summary_df.iterrows():
            organ = row.get("Organ", "Unknown")
            n_samples = row.get("n", 0)
            events = row.get("events", 0)
            
            if not np.isnan(events) and n_samples > 0:
                event_rate = (events / n_samples) * 100
                min_events = 5  # Minimum positive events for stable analysis
                
                validation_checks.append({
                    'Check': f'{organ} Event Rate',
                    'Expected': f'>= {min_events} events for stable analysis',
                    'Actual': f'{int(events)} events ({event_rate:.1f}%)',
                    'Status': 'PASS' if events >= min_events else 'WARNING',
                    'Details': f'{organ}: {int(events)} events out of {n_samples} samples ({event_rate:.1f}% event rate).'
                })
    
    # Check 4: NTCP value range validation
    if combined is not None:
        ntcp_cols = [c for c in combined.columns if c.startswith("NTCP_")]
        for col in ntcp_cols:
            if col in combined.columns:
                ntcp_vals = pd.to_numeric(combined[col], errors="coerce")
                min_val = ntcp_vals.min()
                max_val = ntcp_vals.max()
                nan_count = ntcp_vals.isna().sum()
                total_count = len(ntcp_vals)
                
                in_range = (min_val >= 0) and (max_val <= 1)
                
                validation_checks.append({
                    'Check': f'{col} Value Range',
                    'Expected': '[0, 1]',
                    'Actual': f'[{min_val:.3f}, {max_val:.3f}]',
                    'Status': 'PASS' if in_range else 'FAIL',
                    'Details': f'{col}: Range [{min_val:.3f}, {max_val:.3f}]. NaN values: {nan_count}/{total_count}.'
                })
    
    # Check 5: Model performance consistency
    if not summary_df.empty:
        for _, row in summary_df.iterrows():
            organ = row.get("Organ", "Unknown")
            best_auc = row.get("best_auc", np.nan)
            
            if not np.isnan(best_auc):
                # Check for unrealistic performance
                if best_auc >= 0.95:
                    validation_checks.append({
                        'Check': f'{organ} Model Performance',
                        'Expected': 'AUC < 0.95 (realistic range)',
                        'Actual': f'AUC = {best_auc:.3f}',
                        'Status': 'REVIEW',
                        'Details': f'{organ}: Best model AUC = {best_auc:.3f}. Very high AUC may indicate overfitting or data leakage.'
                    })
                elif best_auc < 0.5:
                    validation_checks.append({
                        'Check': f'{organ} Model Performance',
                        'Expected': 'AUC >= 0.5 (better than random)',
                        'Actual': f'AUC = {best_auc:.3f}',
                        'Status': 'WARNING',
                        'Details': f'{organ}: Best model AUC = {best_auc:.3f}. Performance worse than random guessing.'
                    })
    
    # Check 6: Data pipeline consistency - organ counts
    if not summary_df.empty and "Organ" in summary_df.columns:
        organ_counts = summary_df.groupby("Organ")["n"].sum().to_dict()
        total_combinations = summary_df["n"].sum()
        
        validation_checks.append({
            'Check': 'Organ-Patient Combinations',
            'Expected': 'Matches Step 1 processed_dvh.xlsx count',
            'Actual': f'{int(total_combinations)} total combinations',
            'Status': 'REVIEW',
            'Details': f'Total patient-organ combinations: {int(total_combinations)}. Per organ: {organ_counts}. Verify against Step 1 output.'
        })
    
    # Check 7: Missing DVH files detection
    if combined is not None and "PatientID" in combined.columns and "Organ" in combined.columns:
        # Check if we can infer missing files from patient-organ combinations
        patient_organs = combined.groupby("PatientID")["Organ"].apply(list).to_dict()
        total_expected = len(combined)
        
        validation_checks.append({
            'Check': 'DVH File Coverage',
            'Expected': 'All patient-organ combinations have DVH files',
            'Actual': f'{total_expected} combinations processed',
            'Status': 'REVIEW',
            'Details': f'Processed {total_expected} patient-organ combinations. Verify all DVH files from Step 1 were found in Step 3.'
        })
    
    # Check 8: Biological dose metrics presence
    if combined is not None:
        bio_metrics = [c for c in combined.columns if any(x in c.lower() for x in ['eqd2', 'bed', 'geud', 'eud'])]
        if bio_metrics:
            validation_checks.append({
                'Check': 'Biological Dose Metrics',
                'Expected': 'EQD2, BED, gEUD present',
                'Actual': f'{len(bio_metrics)} biological metrics found',
                'Status': 'PASS',
                'Details': f'Found biological dose metrics: {", ".join(bio_metrics[:5])}'
            })
        else:
            validation_checks.append({
                'Check': 'Biological Dose Metrics',
                'Expected': 'EQD2, BED, gEUD present',
                'Actual': 'No biological metrics found',
                'Status': 'WARNING',
                'Details': 'Biological dose metrics (EQD2, BED, gEUD) not found. Verify Step 3 biological dose calculations.'
            })
    
    # Check 9: SHAP analysis results (if available)
    shap_files = [p for p in files if "shap" in str(p).lower() and p.suffix in [".png", ".jpg", ".svg"]]
    if shap_files:
        validation_checks.append({
            'Check': 'SHAP Explainability',
            'Expected': 'SHAP plots generated for ML models',
            'Actual': f'{len(shap_files)} SHAP plot files found',
            'Status': 'PASS',
            'Details': f'SHAP explainability analysis completed. Found {len(shap_files)} plot files.'
        })
    elif any("ML" in str(c) for c in (combined.columns if combined is not None else [])):
        validation_checks.append({
            'Check': 'SHAP Explainability',
            'Expected': 'SHAP plots generated for ML models',
            'Actual': 'No SHAP plots found',
            'Status': 'INFO',
            'Details': 'ML models present but SHAP plots not found. SHAP analysis may not have been enabled.'
        })
    
    # Create validation DataFrame
    validation_df = pd.DataFrame(validation_checks) if validation_checks else pd.DataFrame()
    
    # ✓ NEW: Add ML validation checks
    print("\n[QA] Checking ML model quality...")
    ml_qa_df = check_ml_model_quality(combined, root)
    
    # Save tables
    excel_path = report_dir / "qa_summary_tables.xlsx"
    with pd.ExcelWriter(excel_path, engine="openpyxl", mode='w') as writer:
        if not summary_df.empty:
            summary_df.to_excel(writer, index=False, sheet_name="PerOrganSummary")
        if patient_ids:
            pd.DataFrame(sorted(list(patient_ids)), columns=["PatientID"]).to_excel(writer, index=False, sheet_name="UniquePatients")
        if not validation_df.empty:
            validation_df.to_excel(writer, index=False, sheet_name="ValidationChecks")
        if not ml_qa_df.empty:
            ml_qa_df.to_excel(writer, index=False, sheet_name="ML_Validation")
    
    print(f"[OK] ML validation checks: {len(ml_qa_df)} checks performed")

    # Save DOCX report
    docx_path = report_dir / "comprehensive_report.docx"
    if Document is None:
        print(f"[WARN] python-docx missing; skipping DOCX. Tables saved at: {excel_path}")
    else:
        doc = Document()
        doc.add_heading("Comprehensive QA Report – NTCP Outputs", 0)
        p = doc.add_paragraph()
        p.add_run("Generated on: ").bold = True
        p.add_run(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"))

        doc.add_heading("1) Summary & Data Integrity", level=1)
        doc.add_paragraph(f"• Estimated unique patients: {global_patients}")
        doc.add_paragraph(f"• Total rows across organs (sum of per‑organ n): {global_rows}")
        doc.add_paragraph("Note: Counting per‑organ files as patients inflates totals. This report estimates unique patients using PatientID (if available) or filename patterns.")

        if not summary_df.empty:
            doc.add_heading("2) Per‑Organ Metrics", level=1)
            show_cols = ["Organ","n","events","event_rate_%","best_model","best_auc"]
            show_cols += [c for c in sorted(summary_df.columns) if c.startswith("AUC|")]
            show_cols += [c for c in sorted(summary_df.columns) if c.startswith("Brier|")]
            table = doc.add_table(rows=1, cols=len(show_cols))
            hdr = table.rows[0].cells
            for i, c in enumerate(show_cols):
                hdr[i].text = c
            for _, row in summary_df[show_cols].iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(show_cols):
                    val = row[c]
                    cells[i].text = "" if (isinstance(val, float) and np.isnan(val)) else str(val)
        else:
            doc.add_paragraph("No per‑organ metrics could be constructed from the provided outputs.")

        doc.add_heading("3) Detected Inconsistencies & Risk Flags", level=1)
        if issues:
            for it in issues:
                doc.add_paragraph(f"• {it}")
        else:
            doc.add_paragraph("No critical issues detected by the heuristic checks.")

        if not summary_df.empty:
            doc.add_heading("4) Model‑Specific Flags by Organ", level=1)
            flags_cols = [c for c in summary_df.columns if c.startswith("Flags|")]
            for _, r in summary_df.iterrows():
                sub_flags = []
                for c in flags_cols:
                    txt = r[c]
                    if isinstance(txt, str) and txt.strip():
                        sub_flags.append(f"{c.split('|',1)[1]} → {txt}")
                if sub_flags:
                    doc.add_paragraph(f"{r['Organ']}:")
                    for f in sub_flags:
                        doc.add_paragraph(f"• {f}")

        doc.add_heading("5) ML Overfitting/Leakage Heuristics (Applied)", level=1)
        doc.add_paragraph("Flagged when AUC ≥ 0.90 with n < 40 or events < 8. High discrimination in small/low‑event cohorts suggests optimism or leakage.")
        
        # Add validation checks section
        if not validation_df.empty:
            doc.add_heading("6) Comprehensive Validation Checks", level=1)
            doc.add_paragraph("Detailed validation checks performed on data pipeline consistency, sample sizes, model performance, and biological dose metrics.")
            
            # Add validation summary
            if not validation_df.empty:
                status_counts = validation_df['Status'].value_counts().to_dict()
                doc.add_paragraph("Validation Status Summary:")
                for status, count in status_counts.items():
                    doc.add_paragraph(f"  • {status}: {count}")
            
            # Add validation table (first 20 rows to avoid overwhelming the document)
            val_table = doc.add_table(rows=1, cols=5)
            val_hdr = val_table.rows[0].cells
            val_hdr[0].text = "Check"
            val_hdr[1].text = "Expected"
            val_hdr[2].text = "Actual"
            val_hdr[3].text = "Status"
            val_hdr[4].text = "Details"
            
            for _, row in validation_df.head(20).iterrows():  # Limit to first 20 for readability
                cells = val_table.add_row().cells
                cells[0].text = str(row.get('Check', ''))
                cells[1].text = str(row.get('Expected', ''))
                cells[2].text = str(row.get('Actual', ''))
                cells[3].text = str(row.get('Status', ''))
                details = str(row.get('Details', ''))
                cells[4].text = details[:80] + "..." if len(details) > 80 else details  # Truncate long details
            
            if len(validation_df) > 20:
                doc.add_paragraph(f"Note: Showing first 20 of {len(validation_df)} validation checks. See Excel file for complete list.")

        doc.save(docx_path)

    print(f"[OK] Saved report: {docx_path}")
    print(f"[OK] Saved tables: {excel_path}")
    
    # ✓ VERIFY files were created
    if docx_path.exists() and excel_path.exists():
        print(f"\n{'='*70}")
        print(f"[OK] QA report complete!")
        print(f"[OK] Report: {docx_path}")
        print(f"[OK] Tables: {excel_path}")
        print(f"{'='*70}\n")
    else:
        print(f"\n[!] Warning: Some QA report files not created!")
        if not docx_path.exists():
            print(f"[!] Missing: {docx_path}")
        if not excel_path.exists():
            print(f"[!] Missing: {excel_path}")
    
    # Print summary statistics
    print("\n" + "=" * 70)
    print("QA REPORT SUMMARY")
    print("=" * 70)
    print(f"Unique Patients: {global_patients}")
    print(f"Total Patient-Organ Combinations: {global_rows}")
    print(f"Organs Analyzed: {len(summary_df) if not summary_df.empty else 0}")
    print(f"Validation Checks: {len(validation_checks)}")
    if not validation_df.empty:
        pass_count = len(validation_df[validation_df['Status'] == 'PASS'])
        warn_count = len(validation_df[validation_df['Status'] == 'WARNING'])
        fail_count = len(validation_df[validation_df['Status'] == 'FAIL'])
        review_count = len(validation_df[validation_df['Status'] == 'REVIEW'])
        print(f"  - PASS: {pass_count}")
        print(f"  - WARNING: {warn_count}")
        print(f"  - FAIL: {fail_count}")
        print(f"  - REVIEW: {review_count}")
    print("=" * 70)


if __name__ == "__main__":
    main()
